"""Ingestion pipeline for GenAI Tutor knowledge sources."""

from __future__ import annotations

import asyncio
import hashlib
import logging
import os
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from uuid import uuid4

from pytesseract import TesseractNotFoundError

from src.backend.services.pdf_ocr_service import pdf_ocr_service
from src.backend.services.tutor_storage_service import tutor_storage_service
from src.genai_tutor.code_analyzer.common_analyzer import extract_code_graph
from src.llm.model_manager import model_manager
from src.rag.vector_store_pg import get_pg_vector_store

logger = logging.getLogger(__name__)

TEXT_FILE_EXTENSIONS = {'.txt', '.md', '.markdown'}
CODE_FILE_EXTENSIONS = {'.py', '.c', '.h'}
IMAGE_FILE_EXTENSIONS = {'.png', '.jpg', '.jpeg'}
SUPPORTED_FILE_EXTENSIONS = (
    TEXT_FILE_EXTENSIONS
    | CODE_FILE_EXTENSIONS
    | IMAGE_FILE_EXTENSIONS
    | {'.pdf', '.docx'}
)
JOB_TERMINAL_STATUSES = {'partial_success', 'success', 'failed'}
MAX_PREVIEW_CHUNKS = 10
DEFAULT_TEXT_CHUNK_SIZE = 1200
DEFAULT_CODE_CHUNK_SIZE = 1800
DEFAULT_FILE_CONCURRENCY = 4


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _isoformat(value: Optional[datetime]) -> Optional[str]:
    if value is None:
        return None
    return value.isoformat()


def _sha256_bytes(payload: bytes) -> str:
    return f'sha256:{hashlib.sha256(payload).hexdigest()}'


def _sanitize_token(value: str) -> str:
    sanitized = ''.join(char if char.isalnum() or char in {'-', '_'} else '-' for char in value.lower())
    sanitized = sanitized.strip('-')
    return sanitized or 'unknown'


def _detect_language(file_path: Path, requested_language: Optional[str]) -> str:
    if requested_language:
        return requested_language.lower()

    extension_map = {
        '.py': 'python',
        '.c': 'c',
        '.h': 'c',
    }
    return extension_map.get(file_path.suffix.lower(), 'mixed')


def _split_text_chunks(content: str, max_chars: int) -> list[str]:
    normalized = content.replace('\r\n', '\n').strip()
    if not normalized:
        return []

    paragraphs = [paragraph.strip() for paragraph in normalized.split('\n\n') if paragraph.strip()]
    chunks: list[str] = []
    current = ''

    for paragraph in paragraphs:
        candidate = paragraph if not current else f'{current}\n\n{paragraph}'
        if len(candidate) <= max_chars:
            current = candidate
            continue

        if current:
            chunks.append(current)
            current = ''

        if len(paragraph) <= max_chars:
            current = paragraph
            continue

        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        line_buffer = ''
        for line in lines:
            line_candidate = line if not line_buffer else f'{line_buffer}\n{line}'
            if len(line_candidate) <= max_chars:
                line_buffer = line_candidate
                continue

            if line_buffer:
                chunks.append(line_buffer)
            line_buffer = line[:max_chars]
            remaining = line[max_chars:]
            while remaining:
                chunks.append(remaining[:max_chars])
                remaining = remaining[max_chars:]

        if line_buffer:
            current = line_buffer

    if current:
        chunks.append(current)

    return chunks


def _split_code_chunks(content: str, max_chars: int) -> list[str]:
    normalized = content.replace('\r\n', '\n').strip()
    if not normalized:
        return []

    lines = normalized.splitlines()
    chunks: list[str] = []
    current_lines: list[str] = []
    current_size = 0

    for line in lines:
        is_boundary = (
            line.startswith('def ')
            or line.startswith('class ')
            or line.startswith('int ')
            or line.startswith('void ')
            or line.startswith('char ')
            or line.startswith('float ')
            or line.startswith('double ')
            or line.startswith('#include')
        )
        line_size = len(line) + 1

        if current_lines and (is_boundary and current_size >= max_chars * 0.5 or current_size + line_size > max_chars):
            chunks.append('\n'.join(current_lines).strip())
            current_lines = []
            current_size = 0

        current_lines.append(line)
        current_size += line_size

    if current_lines:
        chunks.append('\n'.join(current_lines).strip())

    return [chunk for chunk in chunks if chunk]


@dataclass
class TutorChunk:
    chunk_id: str
    document_id: str
    chunk_index: int
    content: str
    content_type: str
    language: str
    topic: Optional[str]
    difficulty: Optional[str]
    token_count: int
    checksum: str
    start_offset: int
    end_offset: int

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class TutorDocument:
    document_id: str
    source_path: str
    source_type: str
    checksum: str
    title: str
    language: str
    course_code: str
    topic: Optional[str]
    difficulty: Optional[str]
    license_tag: Optional[str]
    status: str
    chunk_count: int = 0
    error: Optional[str] = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class IngestJobSummary:
    total_files: int = 0
    processed_files: int = 0
    failed_files: int = 0
    skipped_files: int = 0
    total_chunks: int = 0
    total_entities: int = 0
    total_relations: int = 0

    def to_dict(self) -> dict[str, int]:
        return asdict(self)


@dataclass
class IngestJobState:
    job_id: str
    dataset_version: str
    status: str
    source_path: str
    triggered_by: str
    course_code: str
    language: Optional[str]
    topic: Optional[str]
    difficulty: Optional[str]
    reindex_mode: str
    license_tag: Optional[str]
    dry_run: bool
    created_at: datetime = field(default_factory=_utc_now)
    started_at: Optional[datetime] = None
    finished_at: Optional[datetime] = None
    summary: IngestJobSummary = field(default_factory=IngestJobSummary)
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)
    documents: list[TutorDocument] = field(default_factory=list)
    preview_chunks: list[TutorChunk] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            'jobId': self.job_id,
            'datasetVersion': self.dataset_version,
            'status': self.status,
            'sourcePath': self.source_path,
            'triggeredBy': self.triggered_by,
            'courseCode': self.course_code,
            'language': self.language,
            'topic': self.topic,
            'difficulty': self.difficulty,
            'reindexMode': self.reindex_mode,
            'licenseTag': self.license_tag,
            'dryRun': self.dry_run,
            'createdAt': _isoformat(self.created_at),
            'startedAt': _isoformat(self.started_at),
            'finishedAt': _isoformat(self.finished_at),
            'summary': self.summary.to_dict(),
            'warnings': self.warnings,
            'errors': self.errors,
            'documents': [document.to_dict() for document in self.documents],
            'previewChunks': [chunk.to_dict() for chunk in self.preview_chunks],
        }


class TutorIngestionPipeline:
    """Queue-ready ingestion pipeline with an in-memory job store."""

    def __init__(self) -> None:
        self._jobs: dict[str, IngestJobState] = {}
        self._tasks: dict[str, asyncio.Task[Any]] = {}
        self._lock = asyncio.Lock()

    def get_file_concurrency(self) -> int:
        raw = os.getenv('TUTOR_INGEST_FILE_CONCURRENCY')
        if not raw:
            return DEFAULT_FILE_CONCURRENCY
        try:
            return max(1, int(raw))
        except ValueError:
            return DEFAULT_FILE_CONCURRENCY

    def get_allowed_roots(self) -> list[Path]:
        configured_roots = os.getenv('TUTOR_INGEST_ALLOWED_ROOTS')
        if configured_roots:
            roots = [Path(item).expanduser().resolve() for item in configured_roots.split(',') if item.strip()]
            return roots

        current_root = Path.cwd() / 'data-source'
        workspace_root = Path(__file__).resolve().parents[6] / 'data-source'
        roots = []
        for root in (current_root, workspace_root):
            resolved = root.resolve()
            if resolved not in roots:
                roots.append(resolved)
        return roots

    def get_supported_extensions(self) -> list[str]:
        return sorted(SUPPORTED_FILE_EXTENSIONS)

    async def create_job(
        self,
        *,
        source_path: str,
        triggered_by: str,
        course_code: str,
        language: Optional[str],
        topic: Optional[str],
        difficulty: Optional[str],
        reindex_mode: str,
        license_tag: Optional[str],
        dry_run: bool,
    ) -> IngestJobState:
        resolved_source = self._resolve_source_path(source_path)
        job_id = f'ingest_{_utc_now().strftime("%Y%m%d_%H%M%S")}_{uuid4().hex[:6]}'
        dataset_version = self._build_dataset_version(course_code, language)
        job = IngestJobState(
            job_id=job_id,
            dataset_version=dataset_version,
            status='queued',
            source_path=str(resolved_source),
            triggered_by=triggered_by,
            course_code=course_code,
            language=language.lower() if language else None,
            topic=topic,
            difficulty=difficulty,
            reindex_mode=reindex_mode,
            license_tag=license_tag,
            dry_run=dry_run,
        )

        async with self._lock:
            self._jobs[job_id] = job
            self._tasks[job_id] = asyncio.create_task(self._run_job(job_id, resolved_source))

        await tutor_storage_service.create_job(job.to_dict())

        return job

    async def get_job(self, job_id: str) -> Optional[IngestJobState]:
        async with self._lock:
            job = self._jobs.get(job_id)
        if job is not None:
            return job

        persisted = await tutor_storage_service.fetch_job(job_id)
        if persisted is None:
            return None
        return self._job_from_dict(persisted)

    async def list_jobs(self) -> list[IngestJobState]:
        persisted_jobs = await tutor_storage_service.list_jobs(limit=20)
        if persisted_jobs:
            return [self._job_from_dict(item) for item in persisted_jobs]

        async with self._lock:
            return sorted(self._jobs.values(), key=lambda item: item.created_at, reverse=True)

    async def get_health(self) -> dict[str, Any]:
        async with self._lock:
            active_jobs = [job_id for job_id, task in self._tasks.items() if not task.done()]
            completed_jobs = [job_id for job_id, job in self._jobs.items() if job.status in JOB_TERMINAL_STATUSES]

        return {
            'status': 'ok',
            'service': 'genai-tutor-ingestion',
            'allowedRoots': [str(path) for path in self.get_allowed_roots()],
            'supportedExtensions': self.get_supported_extensions(),
            'activeJobs': active_jobs,
            'completedJobs': completed_jobs[:20],
        }

    def _resolve_source_path(self, source_path: str) -> Path:
        candidate = Path(source_path).expanduser()
        if not candidate.is_absolute():
            candidate = (Path.cwd() / candidate).resolve()
        else:
            candidate = candidate.resolve()

        for root in self.get_allowed_roots():
            try:
                candidate.relative_to(root)
                if candidate.exists():
                    return candidate
            except ValueError:
                continue

        root_list = ', '.join(str(root) for root in self.get_allowed_roots())
        raise ValueError(f'sourcePath must exist under an allowed root: {root_list}')

    def _build_dataset_version(self, course_code: str, language: Optional[str]) -> str:
        course_token = _sanitize_token(course_code)
        language_token = _sanitize_token(language or 'mixed')
        timestamp = _utc_now().strftime('%Y%m%d-%H%M%S')
        return f'{course_token}-{language_token}-{timestamp}'

    async def _run_job(self, job_id: str, resolved_source: Path) -> None:
        job = self._jobs[job_id]
        job.status = 'running'
        job.started_at = _utc_now()

        try:
            files = self._discover_files(resolved_source)
            job.summary.total_files = len(files)
            await tutor_storage_service.update_job(job.to_dict())

            if not files:
                job.warnings.append('No supported files were found under sourcePath')
                job.status = 'success'
                return

            file_semaphore = asyncio.Semaphore(self.get_file_concurrency())
            job_state_lock = asyncio.Lock()

            async def process_with_limit(file_path: Path) -> None:
                async with file_semaphore:
                    await self._process_file(job, file_path, job_state_lock)

            await asyncio.gather(*(process_with_limit(file_path) for file_path in files))

            if job.summary.failed_files and job.summary.processed_files:
                job.status = 'partial_success'
            elif job.summary.failed_files:
                job.status = 'failed'
            else:
                job.status = 'success'
        except Exception as exc:
            logger.exception('Tutor ingest job failed: %s', job_id)
            job.errors.append(str(exc))
            job.status = 'failed'
        finally:
            job.finished_at = _utc_now()
            await tutor_storage_service.update_job(job.to_dict())

    def _discover_files(self, source_path: Path) -> list[Path]:
        if source_path.is_file():
            return [source_path] if source_path.suffix.lower() in SUPPORTED_FILE_EXTENSIONS else []

        files = [
            path for path in sorted(source_path.rglob('*'))
            if path.is_file() and path.suffix.lower() in SUPPORTED_FILE_EXTENSIONS
        ]
        return files

    async def _process_file(
        self,
        job: IngestJobState,
        file_path: Path,
        job_state_lock: asyncio.Lock,
    ) -> None:
        source_type = file_path.suffix.lower().lstrip('.')
        document = TutorDocument(
            document_id=f'doc_{uuid4().hex[:12]}',
            source_path=str(file_path),
            source_type=source_type,
            checksum='',
            title=file_path.name,
            language=_detect_language(file_path, job.language),
            course_code=job.course_code,
            topic=job.topic,
            difficulty=job.difficulty,
            license_tag=job.license_tag,
            status='processing',
        )
        async with job_state_lock:
            job.documents.append(document)

        try:
            content_bytes, extracted_text, content_type = await self._extract_content(file_path)
            document.checksum = _sha256_bytes(content_bytes)

            if not extracted_text.strip():
                document.status = 'skipped'
                document.error = 'No text extracted'
                async with job_state_lock:
                    job.summary.skipped_files += 1
                    await tutor_storage_service.update_job(job.to_dict())
                return

            chunks = self._chunk_document(
                document=document,
                extracted_text=extracted_text,
                content_type=content_type,
                topic=job.topic,
                difficulty=job.difficulty,
            )
            embedding_model = model_manager.get_embedding_info()['id']
            embeddings = []
            if chunks and not job.dry_run:
                embeddings = await get_pg_vector_store().create_embeddings_batch(
                    [chunk.content for chunk in chunks],
                    task_type='retrieval_document',
                )
            document.chunk_count = len(chunks)
            document.status = 'completed'
            await tutor_storage_service.upsert_document(
                document.to_dict(),
                job_id=job.job_id,
                dataset_version=job.dataset_version,
            )
            if not job.dry_run:
                await tutor_storage_service.replace_document_chunks(
                    job_id=job.job_id,
                    dataset_version=job.dataset_version,
                    document_id=document.document_id,
                    embedding_model=embedding_model,
                    chunks=[chunk.to_dict() for chunk in chunks],
                    embeddings=embeddings,
                )
                for chunk in chunks:
                    entities, relations = extract_code_graph(chunk.content, chunk.language)
                    if not entities and not relations:
                        continue
                    await tutor_storage_service.replace_chunk_graph(
                        dataset_version=job.dataset_version,
                        document_id=document.document_id,
                        chunk_id=chunk.chunk_id,
                        entities=[
                            {
                                'entityId': f'entity_{uuid4().hex[:12]}',
                                'entityType': entity.entity_type,
                                'name': entity.name,
                                'canonicalName': entity.canonical_name,
                                'language': entity.language,
                                'properties': entity.properties,
                            }
                            for entity in entities
                        ],
                        relations=[
                            {
                                'relationId': f'rel_{uuid4().hex[:12]}',
                                'relationType': relation.relation_type,
                                'fromCanonicalName': relation.from_name,
                                'toCanonicalName': relation.to_name,
                                'weight': relation.weight,
                            }
                            for relation in relations
                        ],
                    )
            async with job_state_lock:
                job.summary.processed_files += 1
                job.summary.total_chunks += len(chunks)
                job.summary.total_entities += self._estimate_entities(chunks, document.language)
                job.summary.total_relations += self._estimate_relations(chunks, document.language)
                if len(job.preview_chunks) < MAX_PREVIEW_CHUNKS:
                    available_slots = MAX_PREVIEW_CHUNKS - len(job.preview_chunks)
                    job.preview_chunks.extend(chunks[:available_slots])
                await tutor_storage_service.update_job(job.to_dict())
        except Exception as exc:
            logger.exception('Failed to ingest file %s', file_path)
            document.status = 'failed'
            document.error = str(exc)
            await tutor_storage_service.upsert_document(
                document.to_dict(),
                job_id=job.job_id,
                dataset_version=job.dataset_version,
            )
            async with job_state_lock:
                job.summary.failed_files += 1
                job.errors.append(f'{file_path}: {exc}')
                await tutor_storage_service.update_job(job.to_dict())

    async def _extract_content(self, file_path: Path) -> tuple[bytes, str, str]:
        file_bytes = file_path.read_bytes()
        suffix = file_path.suffix.lower()

        if suffix in TEXT_FILE_EXTENSIONS | CODE_FILE_EXTENSIONS:
            return file_bytes, file_bytes.decode('utf-8', errors='ignore'), 'code' if suffix in CODE_FILE_EXTENSIONS else 'text'

        if suffix == '.docx':
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:
                raise RuntimeError('python-docx dependency is not available') from exc

            document = DocxDocument(str(file_path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            return file_bytes, '\n\n'.join(paragraphs), 'text'

        if suffix == '.pdf':
            extracted = await asyncio.to_thread(pdf_ocr_service.extract_text_from_pdf, file_bytes)
            return file_bytes, extracted, 'text'

        if suffix in IMAGE_FILE_EXTENSIONS:
            try:
                import pytesseract
                from PIL import Image
            except ImportError as exc:
                raise RuntimeError('Image OCR dependencies are not available') from exc

            def run_ocr() -> str:
                image = Image.open(file_path)
                return pytesseract.image_to_string(image, lang='eng+vie')

            try:
                extracted = await asyncio.to_thread(run_ocr)
            except TesseractNotFoundError:
                logger.warning(
                    'Skipping image OCR for %s because tesseract is not installed',
                    file_path,
                )
                return file_bytes, '', 'text'
            return file_bytes, extracted, 'text'

        raise ValueError(f'Unsupported file extension: {suffix}')

    def _chunk_document(
        self,
        *,
        document: TutorDocument,
        extracted_text: str,
        content_type: str,
        topic: Optional[str],
        difficulty: Optional[str],
    ) -> list[TutorChunk]:
        if content_type == 'code':
            chunk_contents = _split_code_chunks(extracted_text, DEFAULT_CODE_CHUNK_SIZE)
        else:
            chunk_contents = _split_text_chunks(extracted_text, DEFAULT_TEXT_CHUNK_SIZE)

        chunks: list[TutorChunk] = []
        cursor = 0
        for index, chunk_content in enumerate(chunk_contents):
            start_offset = extracted_text.find(chunk_content, cursor)
            if start_offset < 0:
                start_offset = cursor
            end_offset = start_offset + len(chunk_content)
            cursor = end_offset

            chunks.append(
                TutorChunk(
                    chunk_id=f'chk_{uuid4().hex[:12]}',
                    document_id=document.document_id,
                    chunk_index=index,
                    content=chunk_content,
                    content_type=content_type,
                    language=document.language,
                    topic=topic,
                    difficulty=difficulty,
                    token_count=max(1, len(chunk_content.split())),
                    checksum=_sha256_bytes(chunk_content.encode('utf-8', errors='ignore')),
                    start_offset=start_offset,
                    end_offset=end_offset,
                )
            )

        return chunks

    def _estimate_entities(self, chunks: list[TutorChunk], language: str) -> int:
        if language in {'python', 'c'}:
            return sum(chunk.content.count('def ') + chunk.content.count('class ') + chunk.content.count('#include') for chunk in chunks)
        return max(0, len(chunks) // 2)

    def _estimate_relations(self, chunks: list[TutorChunk], language: str) -> int:
        if language in {'python', 'c'}:
            return max(0, self._estimate_entities(chunks, language) - 1)
        return max(0, len(chunks) - 1)

    def _job_from_dict(self, payload: dict[str, Any]) -> IngestJobState:
        job = IngestJobState(
            job_id=payload['jobId'],
            dataset_version=payload['datasetVersion'],
            status=payload['status'],
            source_path=payload['sourcePath'],
            triggered_by=payload['triggeredBy'],
            course_code=payload['courseCode'],
            language=payload.get('language'),
            topic=payload.get('topic'),
            difficulty=payload.get('difficulty'),
            reindex_mode=payload['reindexMode'],
            license_tag=payload.get('licenseTag'),
            dry_run=payload['dryRun'],
            created_at=datetime.fromisoformat(payload['createdAt']) if payload.get('createdAt') else _utc_now(),
            started_at=datetime.fromisoformat(payload['startedAt']) if payload.get('startedAt') else None,
            finished_at=datetime.fromisoformat(payload['finishedAt']) if payload.get('finishedAt') else None,
            summary=IngestJobSummary(**payload.get('summary', {})),
            warnings=payload.get('warnings', []),
            errors=payload.get('errors', []),
            documents=[TutorDocument(**item) for item in payload.get('documents', [])],
            preview_chunks=[TutorChunk(**item) for item in payload.get('previewChunks', [])],
        )
        return job


tutor_ingestion_pipeline = TutorIngestionPipeline()
